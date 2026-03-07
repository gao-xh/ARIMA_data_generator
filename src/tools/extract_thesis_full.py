from docx import Document
import sys
import os

def extract_text(file_path, output_path):
    print(f"Reading {file_path}")
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append(" | ".join(row_text))

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
        print(f"Text extracted to {output_path}")
        print(f"Total lines: {len(full_text)}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    file_path = r"c:\Users\16179\Desktop\ARIMA\ref\doc\初稿中的新.docx"
    output_path = r"c:\Users\16179\Desktop\ARIMA\ref\doc\thesis_content_extracted.txt"
    extract_text(file_path, output_path)
